from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import tempfile
import os
import re
import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="数学文档解析微服务",
    description="专门解析包含数学公式的Word文档，转换为LaTeX格式",
    version="1.0.0"
)

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class DocxMathParser:
    """Word文档数学公式解析器"""
    
    def __init__(self):
        self.math_ml_to_latex = {
            'mi': lambda x: x.text if x.text else '',
            'mn': lambda x: x.text if x.text else '',
            'mo': lambda x: self._convert_operator(x.text) if x.text else '',
            'mfrac': lambda x: f"\\frac{{{self._parse_element(x[0])}}}{{{self._parse_element(x[1])}}}",
            'msup': lambda x: f"{self._parse_element(x[0])}^{{{self._parse_element(x[1])}}}",
            'msub': lambda x: f"{self._parse_element(x[0])}_{{{self._parse_element(x[1])}}}",
            'msubsup': lambda x: f"{self._parse_element(x[0])}_{{{self._parse_element(x[1])}}}^{{{self._parse_element(x[2])}}}",
            'mroot': lambda x: f"\\sqrt[{self._parse_element(x[1])}]{{{self._parse_element(x[0])}}}",
            'msqrt': lambda x: f"\\sqrt{{{self._parse_element(x[0])}}}",
            'mrow': lambda x: ''.join([self._parse_element(child) for child in x]),
            'mtext': lambda x: x.text if x.text else '',
            'mspace': lambda x: ' ',
        }
    
    def _convert_operator(self, op: str) -> str:
        """转换数学运算符"""
        operator_map = {
            '×': '\\times',
            '÷': '\\div',
            '±': '\\pm',
            '∓': '\\mp',
            '≤': '\\leq',
            '≥': '\\geq',
            '≠': '\\neq',
            '≈': '\\approx',
            '∞': '\\infty',
            '∑': '\\sum',
            '∏': '\\prod',
            '∫': '\\int',
            '∂': '\\partial',
            '∇': '\\nabla',
            'α': '\\alpha',
            'β': '\\beta',
            'γ': '\\gamma',
            'δ': '\\delta',
            'θ': '\\theta',
            'λ': '\\lambda',
            'μ': '\\mu',
            'π': '\\pi',
            'σ': '\\sigma',
            'φ': '\\phi',
            'ω': '\\omega',
        }
        return operator_map.get(op, op)
    
    def _parse_element(self, element) -> str:
        """递归解析MathML元素"""
        if element is None:
            return ''
        
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag in self.math_ml_to_latex:
            return self.math_ml_to_latex[tag](element)
        else:
            # 处理未知标签，尝试获取文本内容
            text = element.text if element.text else ''
            for child in element:
                text += self._parse_element(child)
            return text
    
    def extract_math_from_docx(self, docx_path: str) -> str:
        """从docx文件中提取数学公式和文本"""
        try:
            # 打开docx文件
            doc = docx.Document(docx_path)
            result_text = ""
            
            # 解析段落
            for paragraph in doc.paragraphs:
                para_text = ""
                
                # 处理段落中的runs
                for run in paragraph.runs:
                    if run.text:
                        para_text += run.text
                
                # 检查段落中的数学公式
                math_formulas = self._extract_math_from_paragraph(paragraph)
                for formula in math_formulas:
                    para_text = para_text.replace("[MATH]", f"${formula}$", 1)
                
                if para_text.strip():
                    result_text += para_text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text += cell_text + " | "
                    if row_text:
                        result_text += row_text.rstrip(" | ") + "\n"
            
            # 尝试从XML中提取OLE数学对象
            ole_formulas = self._extract_ole_math(docx_path)
            for formula in ole_formulas:
                result_text += f"\n${formula}$\n"
            
            return result_text.strip()
            
        except Exception as e:
            logger.error(f"解析docx文件时出错: {str(e)}")
            raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")
    
    def _extract_math_from_paragraph(self, paragraph) -> List[str]:
        """从段落中提取数学公式"""
        formulas = []
        try:
            # 检查段落的XML
            p_xml = paragraph._element
            
            # 查找MathML元素
            math_elements = p_xml.xpath('.//m:oMath', namespaces={
                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            })
            
            for math_elem in math_elements:
                try:
                    latex_formula = self._convert_omath_to_latex(math_elem)
                    if latex_formula:
                        formulas.append(latex_formula)
                except Exception as e:
                    logger.warning(f"转换数学公式时出错: {str(e)}")
                    continue
                    
        except Exception as e:
            logger.warning(f"提取段落数学公式时出错: {str(e)}")
        
        return formulas
    
    def _convert_omath_to_latex(self, omath_element) -> str:
        """将Office Math转换为LaTeX"""
        try:
            # 这是一个简化的转换器，实际项目中可能需要更复杂的处理
            text_content = ''.join(omath_element.itertext())
            
            # 基本的转换规则
            latex_text = text_content
            latex_text = re.sub(r'(\w+)_(\w+)', r'\1_{\2}', latex_text)  # 下标
            latex_text = re.sub(r'(\w+)\^(\w+)', r'\1^{\2}', latex_text)  # 上标
            latex_text = re.sub(r'sqrt\(([^)]+)\)', r'\\sqrt{\1}', latex_text)  # 平方根
            
            return latex_text
            
        except Exception as e:
            logger.warning(f"转换Office Math时出错: {str(e)}")
            return ""
    
    def _extract_ole_math(self, docx_path: str) -> List[str]:
        """从docx文件中提取OLE数学对象"""
        formulas = []
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找嵌入的数学对象
                for file_name in zip_file.namelist():
                    if 'embeddings' in file_name and file_name.endswith('.bin'):
                        try:
                            # 这里可以添加更复杂的OLE对象解析逻辑
                            # 目前只是一个占位符
                            pass
                        except Exception as e:
                            logger.warning(f"解析OLE对象时出错: {str(e)}")
                            continue
                            
        except Exception as e:
            logger.warning(f"提取OLE数学对象时出错: {str(e)}")
        
        return formulas

# 全局解析器实例
parser = DocxMathParser()

@app.get("/")
async def root():
    """健康检查端点"""
    return {
        "message": "数学文档解析微服务正在运行",
        "version": "1.0.0",
        "status": "healthy"
    }

@app.post("/parse-docx")
async def parse_docx(file: UploadFile = File(...)):
    """
    解析包含数学公式的Word文档
    
    Args:
        file: 上传的.docx文件
    
    Returns:
        解析后的文本内容，包含LaTeX格式的数学公式
    """
    
    # 验证文件类型
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="只支持.docx格式的文件")
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        try:
            # 保存上传的文件
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
            
            logger.info(f"开始解析文件: {file.filename}")
            
            # 解析文档
            parsed_content = parser.extract_math_from_docx(tmp_path)
            
            logger.info(f"文件解析完成: {file.filename}")
            
            return {
                "success": True,
                "filename": file.filename,
                "content": parsed_content,
                "content_length": len(parsed_content)
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"解析文件时发生未知错误: {str(e)}")
            raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")
        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass

@app.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "healthy", "service": "docx-parser"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001, reload=True) 