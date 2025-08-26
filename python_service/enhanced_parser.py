#!/usr/bin/env python3
"""
增强的Word文档解析器
支持OLE对象、图片、数学公式的提取和处理
"""

import docx
from docx.document import Document
from docx.oxml.ns import qn
import zipfile
import xml.etree.ElementTree as ET
import os
import re
import base64
import tempfile
from typing import List, Dict, Any, Optional
import logging

logger = logging.getLogger(__name__)

class EnhancedDocxParser:
    """增强的Word文档解析器"""
    
    def __init__(self):
        self.ole_objects = []
        self.images = []
        self.math_formulas = []
        self.extracted_text = ""
        
        # 数学符号映射
        self.math_symbols = {
            '∑': '\\sum',
            '∏': '\\prod',
            '∫': '\\int',
            '√': '\\sqrt',
            '∞': '\\infty',
            '≤': '\\leq',
            '≥': '\\geq',
            '≠': '\\neq',
            '≈': '\\approx',
            '±': '\\pm',
            '∓': '\\mp',
            '×': '\\times',
            '÷': '\\div',
            'α': '\\alpha',
            'β': '\\beta',
            'γ': '\\gamma',
            'δ': '\\delta',
            'θ': '\\theta',
            'λ': '\\lambda',
            'μ': '\\mu',
            'π': '\\pi',
            'σ': '\\sigma',
            'φ': '\\phi',
            'ω': '\\omega',
        }
    
    def parse_document(self, docx_path: str) -> Dict[str, Any]:
        """解析Word文档的完整内容"""
        try:
            logger.info(f"开始解析文档: {docx_path}")
            
            # 重置状态
            self.ole_objects = []
            self.images = []
            self.math_formulas = []
            self.extracted_text = ""
            
            # 1. 使用python-docx解析基本内容
            doc = docx.Document(docx_path)
            basic_content = self._extract_basic_content(doc)
            
            # 2. 解析ZIP包中的额外内容
            zip_content = self._extract_from_zip(docx_path)
            
            # 3. 提取OLE对象
            ole_content = self._extract_ole_objects(docx_path)
            
            # 4. 提取图片信息
            image_info = self._extract_images(docx_path)
            
            # 5. 合并所有内容
            combined_content = self._combine_content(
                basic_content, zip_content, ole_content, image_info
            )
            
            logger.info(f"解析完成，提取内容长度: {len(combined_content)}")
            
            return {
                'success': True,
                'content': combined_content,
                'metadata': {
                    'ole_objects_count': len(self.ole_objects),
                    'images_count': len(self.images),
                    'math_formulas_count': len(self.math_formulas),
                    'content_length': len(combined_content)
                }
            }
            
        except Exception as e:
            logger.error(f"文档解析失败: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': ''
            }
    
    def _extract_basic_content(self, doc: Document) -> str:
        """提取基本文本内容"""
        content_parts = []
        
        try:
            # 提取段落内容
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    # 检查是否包含数学符号
                    para_text = self._convert_math_symbols(para_text)
                    content_parts.append(para_text)
            
            # 提取表格内容
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            cell_text = self._convert_math_symbols(cell_text)
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(" | ".join(row_content))
                
                if table_content:
                    content_parts.append("\n".join(table_content))
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.warning(f"基本内容提取出错: {str(e)}")
            return ""
    
    def _extract_from_zip(self, docx_path: str) -> str:
        """从ZIP包中提取额外内容"""
        additional_content = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找document.xml
                if 'word/document.xml' in zip_file.namelist():
                    with zip_file.open('word/document.xml') as xml_file:
                        xml_content = xml_file.read()
                        xml_text = self._extract_text_from_xml(xml_content)
                        if xml_text:
                            additional_content.append(xml_text)
                
                # 查找其他相关XML文件
                for file_name in zip_file.namelist():
                    if (file_name.startswith('word/') and 
                        file_name.endswith('.xml') and 
                        'document' not in file_name):
                        try:
                            with zip_file.open(file_name) as xml_file:
                                xml_content = xml_file.read()
                                xml_text = self._extract_text_from_xml(xml_content)
                                if xml_text and len(xml_text) > 10:
                                    additional_content.append(f"[{file_name}]: {xml_text}")
                        except:
                            continue
            
            return "\n".join(additional_content)
            
        except Exception as e:
            logger.warning(f"ZIP内容提取出错: {str(e)}")
            return ""
    
    def _extract_ole_objects(self, docx_path: str) -> str:
        """提取OLE对象信息"""
        ole_content = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找嵌入对象
                for file_name in zip_file.namelist():
                    if 'embeddings' in file_name:
                        self.ole_objects.append({
                            'name': file_name,
                            'type': 'embedded_object'
                        })
                        ole_content.append(f"[OLE对象: {file_name}]")
                    
                    elif file_name.startswith('word/embeddings/'):
                        try:
                            # 尝试读取OLE对象内容
                            with zip_file.open(file_name) as ole_file:
                                ole_data = ole_file.read()
                                # 这里可以添加更复杂的OLE解析逻辑
                                ole_content.append(f"[数学对象: {file_name}, 大小: {len(ole_data)} bytes]")
                        except:
                            ole_content.append(f"[OLE对象: {file_name} - 无法读取]")
            
            return "\n".join(ole_content)
            
        except Exception as e:
            logger.warning(f"OLE对象提取出错: {str(e)}")
            return ""
    
    def _extract_images(self, docx_path: str) -> str:
        """提取图片信息"""
        image_content = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找图片文件
                for file_name in zip_file.namelist():
                    if (file_name.startswith('word/media/') and 
                        any(file_name.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.emf', '.wmf'])):
                        
                        try:
                            with zip_file.open(file_name) as img_file:
                                img_data = img_file.read()
                                
                                self.images.append({
                                    'name': file_name,
                                    'size': len(img_data),
                                    'type': file_name.split('.')[-1].lower()
                                })
                                
                                # 为图片添加描述性文本
                                img_description = self._generate_image_description(file_name, len(img_data))
                                image_content.append(img_description)
                        except:
                            image_content.append(f"[图片: {file_name} - 无法读取]")
            
            return "\n".join(image_content)
            
        except Exception as e:
            logger.warning(f"图片提取出错: {str(e)}")
            return ""
    
    def _extract_text_from_xml(self, xml_content: bytes) -> str:
        """从XML内容中提取文本"""
        try:
            root = ET.fromstring(xml_content)
            
            # 移除命名空间前缀以简化处理
            for elem in root.iter():
                if '}' in elem.tag:
                    elem.tag = elem.tag.split('}')[1]
            
            # 提取所有文本节点
            text_parts = []
            for text_elem in root.iter('t'):
                if text_elem.text:
                    text_parts.append(text_elem.text)
            
            # 查找数学公式
            for math_elem in root.iter():
                if 'math' in math_elem.tag.lower() or 'equation' in math_elem.tag.lower():
                    math_text = ''.join(math_elem.itertext())
                    if math_text.strip():
                        self.math_formulas.append(math_text.strip())
                        text_parts.append(f"$${math_text.strip()}$$")
            
            return ' '.join(text_parts)
            
        except Exception as e:
            logger.warning(f"XML文本提取出错: {str(e)}")
            return ""
    
    def _convert_math_symbols(self, text: str) -> str:
        """转换数学符号为LaTeX格式"""
        for symbol, latex in self.math_symbols.items():
            text = text.replace(symbol, latex)
        
        # 处理上下标（简单规则）
        text = re.sub(r'([a-zA-Z])_([0-9]+)', r'\1_{\2}', text)  # 下标
        text = re.sub(r'([a-zA-Z])\^([0-9]+)', r'\1^{\2}', text)  # 上标
        
        return text
    
    def _generate_image_description(self, filename: str, size: int) -> str:
        """为图片生成描述性文本"""
        file_type = filename.split('.')[-1].lower()
        
        # 基于文件名和大小推测图片内容
        description_parts = [f"[图片: {os.path.basename(filename)}]"]
        
        if size > 50000:  # 大图片可能是复杂图表
            description_parts.append("[可能包含: 复杂图表、几何图形或详细示意图]")
        elif size > 10000:  # 中等图片可能是公式或简单图表
            description_parts.append("[可能包含: 数学公式、简单图表或示例图]")
        else:  # 小图片可能是符号或简单图标
            description_parts.append("[可能包含: 数学符号、小图标或简单标记]")
        
        # 基于文件名猜测内容
        filename_lower = filename.lower()
        if any(keyword in filename_lower for keyword in ['graph', 'chart', '图', '表']):
            description_parts.append("[推测类型: 图表类]")
        elif any(keyword in filename_lower for keyword in ['formula', 'equation', '公式', '方程']):
            description_parts.append("[推测类型: 数学公式]")
        elif any(keyword in filename_lower for keyword in ['geometry', 'shape', '几何', '图形']):
            description_parts.append("[推测类型: 几何图形]")
        
        return " ".join(description_parts)
    
    def _combine_content(self, basic: str, zip_content: str, ole_content: str, image_info: str) -> str:
        """合并所有提取的内容"""
        content_parts = []
        
        if basic:
            content_parts.append("=== 文档主要内容 ===")
            content_parts.append(basic)
        
        if ole_content:
            content_parts.append("\n=== OLE对象和数学公式 ===")
            content_parts.append(ole_content)
        
        if image_info:
            content_parts.append("\n=== 图片和图表信息 ===")
            content_parts.append(image_info)
        
        if zip_content and len(zip_content) > 100:
            content_parts.append("\n=== 额外提取内容 ===")
            # 只取前1000字符避免内容过长
            content_parts.append(zip_content[:1000] + "..." if len(zip_content) > 1000 else zip_content)
        
        return "\n\n".join(content_parts) 